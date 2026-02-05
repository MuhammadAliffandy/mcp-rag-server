import os
import re
import sys
import pandas as pd
from typing import List
from PyPDF2 import PdfReader
from docx import Document
from langchain_core.documents import Document as LangChainDocument

# Helper log agar tidak crash
def log_safe(msg):
    try:
        log_dir = "logs"
        os.makedirs(log_dir, exist_ok=True)
        with open(os.path.join(log_dir, "server_debug.log"), "a") as f:
            f.write(f"[Processor] {msg}\n")
    except:
        pass

class DocumentProcessor:
    PATIENT_ID_PATTERN = re.compile(r'(?:ID|Patient)\s*[:#]?\s*(\d+)', re.IGNORECASE)

    @staticmethod
    def process_pdf(file_path: str, doc_type: str) -> List[LangChainDocument]:
        try:
            reader = PdfReader(file_path)
            documents = []
            
            # Deep Summary Generation
            all_text = ""
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    all_text += text + "\n"
                    patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
                    documents.append(LangChainDocument(
                        page_content=f"Source: {os.path.basename(file_path)} (Page {i+1})\n\n{text}",
                        metadata={
                            "source": file_path,
                            "page": i + 1,
                            "patient_ids": ",".join(list(set(patient_ids))),
                            "type": "pdf_page",
                            "doc_type": doc_type
                        }
                    ))
            
            # Structured File Summary
            summary_content = f"""
[DEEP SUMMARY] File: {os.path.basename(file_path)}
Format: PDF
Pages: {len(reader.pages)}
Preview: {all_text[:500]}...
Medical Context: This is a document-based medical record or guideline.
            """.strip()
            
            documents.append(LangChainDocument(
                page_content=summary_content,
                metadata={"source": file_path, "type": "file_summary", "doc_type": doc_type}
            ))
            
            return documents
        except Exception as e:
            log_safe(f"Error processing PDF {file_path}: {e}")
            return []

    @staticmethod
    def process_docx(file_path: str, doc_type: str) -> List[LangChainDocument]:
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            documents = []
            
            full_text = "\n".join(paragraphs)
            
            # Structured File Summary
            summary_content = f"""
[DEEP SUMMARY] File: {os.path.basename(file_path)}
Format: DOCX
Paragraphs: {len(paragraphs)}
Preview: {full_text[:500]}...
            """.strip()
            
            documents.append(LangChainDocument(
                page_content=summary_content,
                metadata={"source": file_path, "type": "file_summary", "doc_type": doc_type}
            ))
            
            current_chunk = []
            current_length = 0
            for i, para in enumerate(paragraphs):
                current_chunk.append(para)
                current_length += len(para)
                
                if current_length > 1000 or i == len(paragraphs) - 1:
                    text = "\n".join(current_chunk)
                    patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
                    documents.append(LangChainDocument(
                        page_content=f"Source: {os.path.basename(file_path)}\n\n{text}",
                        metadata={
                            "source": file_path,
                            "chunk_index": len(documents),
                            "patient_ids": ",".join(list(set(patient_ids))),
                            "type": "docx_chunk",
                            "doc_type": doc_type
                        }
                    ))
                    current_chunk = []
                    current_length = 0
            return documents
        except Exception as e:
            log_safe(f"Error processing DOCX {file_path}: {e}")
            return []

    @staticmethod
    def process_tabular(file_path: str, doc_type: str) -> List[LangChainDocument]:
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
                
            patient_id_cols = [col for col in df.columns if 'id' in col.lower() or 'patient' in col.lower()]
            documents = []
            
            # Deep Tabular Analysis for RAG
            cols_info = []
            for col in df.columns:
                dtype = str(df[col].dtype)
                nulls = df[col].isnull().sum()
                sample = str(df[col].iloc[0]) if not df.empty else "N/A"
                cols_info.append(f"- {col} ({dtype}): {nulls} missing, example: {sample}")
            
            cols_summary = "\n".join(cols_info)
            summary_content = f"""
[DEEP SUMMARY] File: {os.path.basename(file_path)}
Format: Tabular (Excel/CSV)
Rows: {len(df)}, Columns: {len(df.columns)}
Columns Overview:
{cols_summary}

Statistical Highlights:
{df.describe().to_string() if not df.select_dtypes('number').empty else "No numeric columns."}
            """.strip()
            
            documents.append(LangChainDocument(
                page_content=summary_content,
                metadata={
                    "source": file_path,
                    "type": "file_summary",
                    "doc_type": doc_type,
                    "df_json": df.to_json() # Keep for server data sync
                }
            ))
            
            max_rows = 50 
            for i, row in df.head(max_rows).iterrows():
                row_content = [f"{col}: {val}" for col, val in row.items() if pd.notnull(val)]
                row_text = f"Data Record (Row {i+1}): {', '.join(row_content)}"
                p_id = str(row[patient_id_cols[0]]) if patient_id_cols else ""

                documents.append(LangChainDocument(
                    page_content=row_text,
                    metadata={
                        "source": file_path,
                        "patient_ids": p_id,
                        "row_index": i,
                        "type": "tabular_row",
                        "doc_type": doc_type
                    }
                ))
            return documents
        except Exception as e:
            log_safe(f"Error processing Tabular {file_path}: {e}")
            return []

    @staticmethod
    def process_image(file_path: str, doc_type: str) -> List[LangChainDocument]:
        text = f"Image file: {os.path.basename(file_path)}"
        try:
            from PIL import Image
            import pytesseract
            text = pytesseract.image_to_string(Image.open(file_path))
        except:
            pass 
            
        patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
        if not patient_ids:
             patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(file_path)
            
        return [LangChainDocument(
            page_content=f"Source: {os.path.basename(file_path)}\n\n{text}",
            metadata={
                "source": file_path,
                "patient_ids": ",".join(list(set(patient_ids))),
                "type": "image",
                "doc_type": doc_type
            }
        )]

    @classmethod
    def load_directory(cls, directory_path: str, doc_type: str = "internal_patient") -> List[LangChainDocument]:
        all_docs = []
        if not os.path.exists(directory_path):
            return []
            
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            if filename.startswith('.'): continue
            
            ext = filename.lower()
            if ext.endswith('.pdf'):
                all_docs.extend(cls.process_pdf(file_path, doc_type))
            elif ext.endswith('.docx'):
                all_docs.extend(cls.process_docx(file_path, doc_type))
            elif ext.endswith(('.xlsx', '.xls', '.csv')):
                all_docs.extend(cls.process_tabular(file_path, doc_type))
            elif ext.endswith(('.png', '.jpg', '.jpeg')):
                all_docs.extend(cls.process_image(file_path, doc_type))
            elif ext.endswith('.txt'):
                all_docs.extend(cls.process_txt(file_path, doc_type))
                
        return all_docs

    @staticmethod
    def process_txt(file_path: str, doc_type: str) -> List[LangChainDocument]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
            
            summary_content = f"[DEEP SUMMARY] File: {os.path.basename(file_path)}\nFormat: Text\nPreview: {text[:500]}..."
            
            return [
                LangChainDocument(page_content=summary_content, metadata={"source": file_path, "type": "file_summary", "doc_type": doc_type}),
                LangChainDocument(
                    page_content=f"Source: {os.path.basename(file_path)}\n\n{text}",
                    metadata={
                        "source": file_path,
                        "patient_ids": ",".join(list(set(patient_ids))),
                        "type": "text_file",
                        "doc_type": doc_type
                    }
                )
            ]
        except Exception as e:
            log_safe(f"Error processing TXT {file_path}: {e}")
            return []