import os
import re
import pandas as pd
import cv2
from PIL import Image
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
from langchain_core.documents import Document as LangChainDocument

class DocumentProcessor:
    """
    Handles loading and parsing of various file formats: PDF, Word, Excel, CSV, Images.
    Extracts text and metadata (like patient IDs and document types).
    """
    
    PATIENT_ID_PATTERN = re.compile(r'ID\s*(\d+)', re.IGNORECASE)

    @staticmethod
    def process_pdf(file_path: str, doc_type: str = "internal_patient") -> List[LangChainDocument]:
        reader = PdfReader(file_path)
        documents = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
                documents.append(LangChainDocument(
                    page_content=f"Source: {os.path.basename(file_path)}\n\n{text}",
                    metadata={
                        "source": file_path,
                        "page": i + 1,
                        "patient_ids": ",".join(list(set(patient_ids))),
                        "type": "pdf",
                        "doc_type": doc_type
                    }
                ))
        return documents

    @staticmethod
    def process_docx(file_path: str, doc_type: str = "internal_patient") -> List[LangChainDocument]:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = "\n".join(full_text)
        patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
        return [LangChainDocument(
            page_content=f"Source: {os.path.basename(file_path)}\n\n{text}",
            metadata={
                "source": file_path,
                "patient_ids": ",".join(list(set(patient_ids))),
                "type": "docx",
                "doc_type": doc_type
            }
        )]

    @staticmethod
    def process_tabular(file_path: str, doc_type: str = "internal_patient") -> List[LangChainDocument]:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
            
        patient_id_cols = [col for col in df.columns if 'id' in col.lower() or 'patient' in col.lower()]
        all_patient_ids = []
        if patient_id_cols:
            all_patient_ids = df[patient_id_cols[0]].astype(str).tolist()

        content = df.to_string()
        return [LangChainDocument(
            page_content=f"Source: {os.path.basename(file_path)}\n\n{content}",
            metadata={
                "source": file_path,
                "patient_ids": ",".join(list(set(all_patient_ids))),
                "type": "tabular",
                "doc_type": doc_type,
                "df_json": df.to_json() 
            }
        )]

    @staticmethod
    def process_image(file_path: str, doc_type: str = "internal_patient") -> List[LangChainDocument]:
        """
        Processes images using basic OCR or metadata (fallback for now).
        """
        try:
            import pytesseract
            # Try to use OCR
            text = pytesseract.image_to_string(Image.open(file_path))
        except ImportError:
            # Fallback to just the filename if tesseract not found
            text = f"Image file: {os.path.basename(file_path)}"
        
        patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(text)
        if not patient_ids:
            # Try extracting from filename
            patient_ids = DocumentProcessor.PATIENT_ID_PATTERN.findall(file_path)
            
        return [LangChainDocument(
            page_content=f"Source: {os.path.basename(file_path)}\n\n{text}",
            metadata={
                "source": file_path,
                "patient_ids": ",".join(list(set(patient_ids))),
                "type": "image",
                "doc_type": doc_type
            }
        )]

    @classmethod
    def load_directory(cls, directory_path: str, doc_type: str = "internal_patient") -> List[LangChainDocument]:
        all_docs = []
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            ext = filename.lower()
            if ext.endswith('.pdf'):
                all_docs.extend(cls.process_pdf(file_path, doc_type))
            elif ext.endswith('.docx'):
                all_docs.extend(cls.process_docx(file_path, doc_type))
            elif ext.endswith(('.xlsx', '.xls', '.csv')):
                all_docs.extend(cls.process_tabular(file_path, doc_type))
            elif ext.endswith(('.png', '.jpg', '.jpeg')):
                all_docs.extend(cls.process_image(file_path, doc_type))
        return all_docs
